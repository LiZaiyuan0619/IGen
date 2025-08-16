# -*- coding: utf-8 -*-
"""
Markdown到Word文档转换器

目标：
- 解决一级标题居中问题
- 统一中英文字体设置
- 修复目录缩进层级显示
- 优化整体文档格式

作者: Claude
环境: Windows 11, PowerShell

输入: Markdown格式文本内容
执行步骤:
1. 解析Markdown结构（标题、段落、列表、代码块、图片等）
2. 设置统一的文档样式和字体
3. 按层级处理标题并设置正确的对齐方式
4. 处理列表的缩进层级关系
5. 生成格式良好的Word文档

输出: 格式化的Word文档文件
"""

import re
import os
from typing import List, Dict, Tuple, Optional
import tempfile
import zipfile
from functools import lru_cache

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_SECTION
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx 未安装，请运行: pip install python-docx")

# Pandoc / pypandoc 依赖（仅用于公式片段转换为 OMML）
try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False
    print("⚠️ pypandoc 未安装，Pandoc 公式转换不可用。可运行: pip install pypandoc")

# 图片处理依赖
try:
    from PIL import Image
    PIL_AVAILABLE = True
    print("✅ PIL 可用，支持实际图片插入")
except ImportError:
    PIL_AVAILABLE = False
    print("⚠️ PIL 未安装，图片将以文本形式显示。可运行: pip install Pillow")

# LaTeX公式处理不再需要外部依赖
print("✅ 使用Word原生格式处理LaTeX公式")


class MarkdownToWordConverter:
    """高质量的Markdown到Word转换器"""
    
    def __init__(self):
        self.doc = None
        self.current_list_level = 0  # 当前列表层级
        self._first_h1_added = False  # 记录是否已经添加了第一个一级标题
        # 章节与Pandoc状态
        self._current_section_title: Optional[str] = None
        # 全局启用 Pandoc 数学转换
        self._use_pandoc_math: bool = True
        self._pandoc_ready: Optional[bool] = None
        self._pandoc_cache: Dict[Tuple[str, bool], Optional[Tuple[str, str]]] = {}
        # 自然段首行缩进控制（目录之后、参考文献之前；默认启用，目录期间关闭）
        self._indent_paragraphs: bool = True
        self._first_line_indent_pt: int = 24  # 约等于两个汉字宽（12pt 字号）
        self._in_toc: bool = False  # 目录区间内标记
        self.image_stats = {  # 图片处理统计
            'total_found': 0,      # 发现的图片总数
            'inserted': 0,         # 成功插入的图片数
            'fallback': 0          # 使用降级处理的图片数
        }
        self.html_cleanup_stats = {  # HTML清理统计
            'html_blocks_found': 0,    # 发现的HTML块数量
            'html_blocks_removed': 0,  # 删除的HTML块数量
            'characters_removed': 0    # 删除的字符数量
        }
        self.formula_stats = {  # 公式处理统计
            'display_formulas_found': 0,    # 发现的单行公式数量（$$...$$）
            'inline_formulas_found': 0,     # 发现的行内公式数量（$...$）
            'formulas_rendered': 0,         # 成功渲染为图片的公式数量
            'formulas_fallback': 0          # 降级处理的公式数量
        }
        
    def convert(self, markdown_content: str, word_path: str, title: str = "") -> bool:
        """
        转换Markdown内容为Word文档
        
        Args:
            markdown_content: Markdown格式内容
            word_path: Word文档保存路径  
            title: 文档标题
            
        Returns:
            bool: 转换是否成功
        """
        if not DOCX_AVAILABLE:
            print("❌ python-docx未安装，无法生成Word文档")
            return False
            
        try:
            # 创建新文档
            self.doc = Document()
            
            # 设置文档基础样式
            self._setup_document_styles()
            
            # 添加页码
            self._add_page_numbers()
            
            # 添加文档标题
            if title:
                self._add_document_title(title)
            
            # 预处理：清理HTML内容
            cleaned_markdown = self._preprocess_remove_html(markdown_content)
            
            # 预处理：移除所有出现的转义美元符号 \$（包括正文、参考、公式、图片、表格等所有内容）
            cleaned_markdown = self._preprocess_remove_escaped_dollar(cleaned_markdown)
            
            # 预处理：处理LaTeX数学公式
            formula_processed_markdown = self._preprocess_latex_formulas(cleaned_markdown)
            
            # 解析并处理内容（跳过第一个标题如果与文档标题重复）
            self._process_markdown_content(formula_processed_markdown, skip_first_title=bool(title))
            
            # 后处理：修复所有未处理的加粗格式
            self._fix_remaining_bold_formatting()
            
            # 验证所有章节的页眉设置
            self._verify_all_headers()
            
            # 首页封面后处理：标题居中到整页中央并追加生成时间
            try:
                self._postprocess_cover_page()
            except Exception as cover_e:
                print(f"⚠️ 首页后处理失败: {cover_e}")

            # 保存文档
            word_dir = os.path.dirname(word_path)
            if word_dir:  # 只有当有目录路径时才创建目录
                os.makedirs(word_dir, exist_ok=True)
            self.doc.save(word_path)
            print(f"📄 Word文档已保存到: {word_path}")
            
            # 输出图片处理统计信息
            self._print_image_statistics()
            
            return True
            
        except Exception as e:
            print(f"❌ 生成Word文档时出错: {e}")
            return False
    
    def _setup_document_styles(self):
        """设置文档基础样式和字体"""
        
        # 设置正文样式
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'  # 英文字体
        normal_font.size = Pt(12)
        
        # 设置中文字体
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置段落格式
        paragraph_format = normal_style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = 1.5  # 1.5倍行距
        paragraph_format.space_after = Pt(6)  # 段后间距
        
        # 创建自定义标题样式
        self._create_heading_styles()
        
        # 创建自定义列表样式
        self._create_list_styles()
    
    def _create_heading_styles(self):
        """创建自定义标题样式"""
        
        # 一级标题样式（居中，大字体）
        try:
            h1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Times New Roman'
            h1_font.size = Pt(18)
            h1_font.bold = True
            h1_font.color.rgb = RGBColor(0, 0, 0)
            
            # 设置中文字体
            h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            # 设置段落格式 - 居中对齐
            h1_paragraph = h1_style.paragraph_format
            h1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1_paragraph.space_before = Pt(18)
            h1_paragraph.space_after = Pt(12)
        except:
            pass  # 样式可能已存在
            
        # 二级标题样式（左对齐）
        try:
            h2_style = self.doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Times New Roman'
            h2_font.size = Pt(16)
            h2_font.bold = True
            h2_font.color.rgb = RGBColor(0, 0, 0)
            
            # 设置中文字体
            h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            # 设置段落格式
            h2_paragraph = h2_style.paragraph_format
            h2_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h2_paragraph.space_before = Pt(12)
            h2_paragraph.space_after = Pt(6)
            h2_paragraph.left_indent = Pt(0)  # 二级标题不缩进
        except:
            pass
            
        # 三级标题样式
        try:
            h3_style = self.doc.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
            h3_font = h3_style.font
            h3_font.name = 'Times New Roman'
            h3_font.size = Pt(14)
            h3_font.bold = True
            h3_font.color.rgb = RGBColor(0, 0, 0)
            
            # 设置中文字体
            h3_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            # 设置段落格式
            h3_paragraph = h3_style.paragraph_format
            h3_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h3_paragraph.space_before = Pt(6)
            h3_paragraph.space_after = Pt(6)
            # 三级标题左对齐且无缩进，保持与二级标题左边对齐
            h3_paragraph.left_indent = Pt(0)
            try:
                h3_paragraph.first_line_indent = Pt(0)
                h3_paragraph.hanging_indent = Pt(0)
            except Exception:
                pass
        except:
            pass
    
    def _create_list_styles(self):
        """创建自定义列表样式"""
        
        # 一级列表样式
        try:
            list1_style = self.doc.styles.add_style('CustomList1', WD_STYLE_TYPE.PARAGRAPH)
            list1_font = list1_style.font
            list1_font.name = 'Times New Roman'
            list1_font.size = Pt(12)
            
            # 设置中文字体
            list1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 设置段落格式
            list1_paragraph = list1_style.paragraph_format
            list1_paragraph.left_indent = Pt(18)  # 一级缩进
            list1_paragraph.hanging_indent = Pt(18)
            list1_paragraph.space_after = Pt(3)
        except:
            pass
            
        # 二级列表样式
        try:
            list2_style = self.doc.styles.add_style('CustomList2', WD_STYLE_TYPE.PARAGRAPH)
            list2_font = list2_style.font
            list2_font.name = 'Times New Roman'
            list2_font.size = Pt(12)
            
            # 设置中文字体
            list2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 设置段落格式
            list2_paragraph = list2_style.paragraph_format
            list2_paragraph.left_indent = Pt(54)  # 二级缩进增加到54pt
            list2_paragraph.hanging_indent = Pt(18)
            list2_paragraph.space_after = Pt(3)
        except:
            pass
    
    def _add_page_numbers(self):
        """添加页码"""
        try:
            # 获取文档的sections
            section = self.doc.sections[0]
            
            # 在页脚中添加页码
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码字段
            run = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)
            
        except Exception as e:
            print(f"⚠️ 添加页码时出错: {e}")
    
    def _add_document_title(self, title: str):
        """添加文档标题"""
        title_paragraph = self.doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置标题字体
        for run in title_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(20)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    def _set_section_header(self, section, chapter_title: str):
        """
        为指定节设置页眉
        
        Args:
            section: Word文档的节对象
            chapter_title: 章节标题
        """
        try:
            # 获取节的页眉
            header = section.header
            
            # 关键修复：断开与前一节的页眉链接，确保每个节都有独立的页眉
            try:
                header.is_linked_to_previous = False
            except Exception as link_error:
                print(f"⚠️ 断开页眉链接失败: {link_error}")
                # 如果断开链接失败，尝试其他方法
                try:
                    # 尝试通过XML操作断开链接
                    header_part = header._element
                    if header_part is not None:
                        # 移除titlePg属性可能有助于独立设置
                        pass
                except:
                    pass
            
            # 清除页眉中的现有内容
            for paragraph in header.paragraphs:
                paragraph.clear()
            
            # 创建页眉段落或使用第一个段落
            if header.paragraphs:
                header_para = header.paragraphs[0]
            else:
                header_para = header.add_paragraph()
            
            # 截断过长的标题（防止页眉过宽）
            if len(chapter_title) > 80:
                display_title = chapter_title[:77] + "..."
            else:
                display_title = chapter_title
            
            # 设置页眉内容
            run = header_para.add_run(display_title)
            
            # 设置页眉样式
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置字体样式
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = False
            run.font.italic = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 设置页眉与正文的间距
            header_para.paragraph_format.space_after = Pt(6)
            
            # 添加分隔线（可选）
            self._add_header_border(header_para)
            
        except Exception as e:
            print(f"⚠️ 设置页眉时出错: {e}")
            # 如果设置失败，尝试简化版本
            try:
                # 简化版本也要先断开链接
                try:
                    section.header.is_linked_to_previous = False
                    print(f"✅ 简化模式下已断开页眉链接")
                except:
                    pass
                    
                header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
                header_para.text = chapter_title[:50] + ("..." if len(chapter_title) > 50 else "")
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"✅ 使用简化版本设置页眉: {chapter_title[:30]}...")
            except Exception as fallback_error:
                print(f"❌ 简化版本也失败: {fallback_error}")
                pass  # 如果连简化版本都失败，就放弃页眉设置
    
    def _add_header_border(self, paragraph):
        """为页眉段落添加下边框"""
        try:
            # 创建下边框
            from docx.oxml.parser import parse_xml
            border_xml = """
            <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>
            </w:pBdr>
            """
            border_element = parse_xml(border_xml)
            paragraph._p.get_or_add_pPr().append(border_element)
            
        except Exception as e:
            # 如果添加边框失败，静默忽略（非关键功能）
            pass
    
    def _verify_header_independence(self, section, expected_title: str):
        """
        验证页眉是否独立设置成功
        
        Args:
            section: 当前节
            expected_title: 期望的标题
        """
        try:
            header = section.header
            
            # 检查页眉链接状态
            is_linked = getattr(header, 'is_linked_to_previous', None)
            
            # 检查页眉内容
            if header.paragraphs:
                actual_content = header.paragraphs[0].text

                # 检查是否匹配
                if expected_title in actual_content or actual_content in expected_title:
                    pass
                else:
                    print(f"⚠️ 页眉内容不匹配！")
            else:
                print(f"⚠️ 页眉中没有段落")
                
        except Exception as e:
            print(f"⚠️ 验证页眉时出错: {e}")
    
    def _preprocess_remove_html(self, markdown_content: str) -> str:
        """
        预处理：移除Markdown内容中的HTML块
        
        Args:
            markdown_content: 原始Markdown内容
            
        Returns:
            str: 清理后的Markdown内容
        """
        print("🧹 开始HTML内容预处理...")
        
        original_length = len(markdown_content)
        
        # 检测HTML块
        html_blocks = self._detect_html_blocks(markdown_content)
        
        if not html_blocks:
            print("✅ 未发现HTML内容，无需清理")
            return markdown_content
        
        # 清理HTML块
        cleaned_content = self._remove_html_blocks(markdown_content, html_blocks)
        
        # 更新统计信息
        cleaned_length = len(cleaned_content)
        self.html_cleanup_stats['characters_removed'] = original_length - cleaned_length
        
        # 清理统计信息已整合到最终报告中
        
        return cleaned_content
    
    def _detect_html_blocks(self, content: str) -> list:
        """
        检测内容中的HTML块
        
        Args:
            content: 要检测的内容
            
        Returns:
            list: HTML块的列表，每个元素是(start_pos, end_pos, block_content)
        """
        html_blocks = []
        
        # 使用正则表达式匹配HTML块
        # 模式：<html>......</html>，支持多行和贪婪匹配
        html_pattern = r'<html>.*?</html>'
        
        try:
            matches = re.finditer(html_pattern, content, re.DOTALL | re.IGNORECASE)
            
            for match in matches:
                start_pos = match.start()
                end_pos = match.end()
                block_content = match.group(0)
                
                html_blocks.append((start_pos, end_pos, block_content))
                self.html_cleanup_stats['html_blocks_found'] += 1
                
                # 显示找到的HTML块信息（截取前100字符）
                preview = block_content[:100] + "..." if len(block_content) > 100 else block_content

        
        except Exception as e:
            print(f"⚠️ 检测HTML块时出错: {e}")
        
        return html_blocks
    
    def _remove_html_blocks(self, content: str, html_blocks: list) -> str:
        """
        从内容中移除HTML块
        
        Args:
            content: 原始内容
            html_blocks: 要移除的HTML块列表
            
        Returns:
            str: 清理后的内容
        """
        if not html_blocks:
            return content
        
        # 按位置倒序排序，从后往前删除（避免位置偏移）
        html_blocks.sort(key=lambda x: x[0], reverse=True)
        
        cleaned_content = content
        
        for start_pos, end_pos, block_content in html_blocks:
            try:
                # 删除HTML块
                cleaned_content = cleaned_content[:start_pos] + cleaned_content[end_pos:]
                self.html_cleanup_stats['html_blocks_removed'] += 1
                
                
            except Exception as e:
                print(f"⚠️ 删除HTML块时出错: {e}")
        
        # 清理多余的空行（HTML删除后可能留下）
        cleaned_content = self._clean_extra_newlines(cleaned_content)
        
        return cleaned_content
    
    def _clean_extra_newlines(self, content: str) -> str:
        """
        清理多余的换行符
        
        Args:
            content: 要清理的内容
            
        Returns:
            str: 清理后的内容
        """
        # 将多个连续的换行符替换为最多两个换行符
        cleaned = re.sub(r'\n{3,}', '\n\n', content)
        
        # 移除文档末尾的多余换行符
        cleaned = cleaned.rstrip('\n') + '\n'
        
        return cleaned

    def _preprocess_remove_escaped_dollar(self, content: str) -> str:
        """
        预处理：
        - 移除文档中所有转义美元符号（\$）
        - 移除所有非法的双美元符号（$$）。仅当 $$ 后紧跟换行符（\n 或 \r\n）时视为合法并保留（例如："$$\n"）。
        注意：该规则会清理诸如 " $$ ", "$$ 文本", "文本 $$" 等情况；仅保留行尾的 "$$" 紧随换行符的分隔用法。
        作用范围：全文（正文、参考、公式、图片、表格等）。
        """
        try:
            # 1) 删除所有 "\$" 字符序列
            processed = content.replace('\\$', '')

            # 2) 删除所有不直接跟随换行符的 "$$"
            #    仅保留形如 "$$\n" 或 "$$\r\n" 的合法标记
            processed = re.sub(r'\$\$(?!\r?\n)', '', processed)

            return processed
        except Exception:
            return content
    
    def _preprocess_latex_formulas(self, markdown_content: str) -> str:
        """
        预处理：处理LaTeX数学公式，转换为Word友好格式
        
        Args:
            markdown_content: 包含LaTeX公式的Markdown内容
            
        Returns:
            str: 处理后的Markdown内容
        """
        
        # 检测并处理单行公式（$$...$$）
        content_with_display = self._process_display_formulas(markdown_content)
        
        # 检测并处理行内公式（$...$）
        content_with_inline = self._process_inline_formulas(content_with_display)
        
        return content_with_inline
    
    def _process_display_formulas(self, content: str) -> str:
        """
        处理单行公式（$$...$$）
        
        Args:
            content: 要处理的内容
            
        Returns:
            str: 处理后的内容
        """
        # 匹配$$...$$格式的公式（支持多行）
        display_pattern = r'\$\$\s*(.*?)\s*\$\$'
        
        def replace_display_formula(match):
            latex_code = match.group(1).strip()
            if not latex_code:
                return match.group(0)  # 空公式，保持原样
            
            self.formula_stats['display_formulas_found'] += 1
            
            # 清理LaTeX代码
            cleaned_latex = self._clean_latex_code(latex_code)
            
            # 标记为单行公式，稍后在文档中特殊处理
            self.formula_stats['formulas_rendered'] += 1
            return f"\n\n<<<DISPLAY_FORMULA>>>{cleaned_latex}<<<END_DISPLAY_FORMULA>>>\n\n"
        
        try:
            processed_content = re.sub(display_pattern, replace_display_formula, content, flags=re.DOTALL)
            return processed_content
        except Exception as e:
            print(f"⚠️ 处理单行公式时出错: {e}")
            return content
    
    def _process_inline_formulas(self, content: str) -> str:
        """
        处理行内公式（$...$）
        
        Args:
            content: 要处理的内容
            
        Returns:
            str: 处理后的内容
        """
        # 匹配$...$格式的公式，但避免匹配$$...$$
        # 使用负向后顾和负向先行断言
        inline_pattern = r'(?<!\$)\$([^$\n]+?)\$(?!\$)'
        
        def replace_inline_formula(match):
            latex_code = match.group(1).strip()
            if not latex_code:
                return match.group(0)  # 空公式，保持原样
            
            self.formula_stats['inline_formulas_found'] += 1
            
            # 清理LaTeX代码
            cleaned_latex = self._clean_latex_code(latex_code)
            
            # 尝试转换为格式化文本
            formatted_text = self._convert_latex_to_formatted_text(cleaned_latex)
            
            if formatted_text:
                self.formula_stats['formulas_rendered'] += 1
                return f"<<<INLINE_FORMULA>>>{formatted_text}<<<END_INLINE_FORMULA>>>"
            else:
                # 转换失败，使用降级方案
                self.formula_stats['formulas_fallback'] += 1
                return f"<<<INLINE_FORMULA>>>{cleaned_latex}<<<END_INLINE_FORMULA>>>"
        
        try:
            processed_content = re.sub(inline_pattern, replace_inline_formula, content)
            return processed_content
        except Exception as e:
            print(f"⚠️ 处理行内公式时出错: {e}")
            return content
    
    def _clean_latex_code(self, latex_code: str) -> str:
        """
        清理LaTeX代码中的多余空格
        
        Args:
            latex_code: 原始LaTeX代码
            
        Returns:
            str: 清理后的LaTeX代码
        """
        # 移除花括号前后的多余空格
        # 例如: "h _ { q, l }" → "h_{q,l}"
        
        # 处理下标：_ { } → _{}
        latex_code = re.sub(r'_\s*\{\s*(.*?)\s*\}', r'_{\1}', latex_code)
        
        # 处理上标：^ { } → ^{}
        latex_code = re.sub(r'\^\s*\{\s*(.*?)\s*\}', r'^{\1}', latex_code)
        
        # 处理一般花括号：{ } → {}，但保留必要的空格
        latex_code = re.sub(r'\{\s*(.*?)\s*\}', r'{\1}', latex_code)
        
        # 移除命令后的多余空格：\alpha → \alpha
        latex_code = re.sub(r'\\([a-zA-Z]+)\s+', r'\\\1 ', latex_code)
        
        # 移除多余的空格
        latex_code = re.sub(r'\s+', ' ', latex_code).strip()
        
        return latex_code
    
    def _convert_latex_to_formatted_text(self, latex_code: str) -> Optional[str]:
        """
        将简单的LaTeX代码转换为格式化文本
        
        Args:
            latex_code: 清理后的LaTeX代码
            
        Returns:
            str: 格式化文本，如果无法转换则返回None
        """
        # 希腊字母映射
        greek_letters = {
            r'\\alpha': 'α', r'\\beta': 'β', r'\\gamma': 'γ', r'\\delta': 'δ',
            r'\\epsilon': 'ε', r'\\zeta': 'ζ', r'\\eta': 'η', r'\\theta': 'θ',
            r'\\iota': 'ι', r'\\kappa': 'κ', r'\\lambda': 'λ', r'\\mu': 'μ',
            r'\\nu': 'ν', r'\\xi': 'ξ', r'\\pi': 'π', r'\\rho': 'ρ',
            r'\\sigma': 'σ', r'\\tau': 'τ', r'\\upsilon': 'υ', r'\\phi': 'φ',
            r'\\chi': 'χ', r'\\psi': 'ψ', r'\\omega': 'ω',
            # 大写希腊字母
            r'\\Alpha': 'Α', r'\\Beta': 'Β', r'\\Gamma': 'Γ', r'\\Delta': 'Δ',
            r'\\Epsilon': 'Ε', r'\\Zeta': 'Ζ', r'\\Eta': 'Η', r'\\Theta': 'Θ',
            r'\\Lambda': 'Λ', r'\\Mu': 'Μ', r'\\Nu': 'Ν', r'\\Xi': 'Ξ',
            r'\\Pi': 'Π', r'\\Rho': 'Ρ', r'\\Sigma': 'Σ', r'\\Tau': 'Τ',
            r'\\Phi': 'Φ', r'\\Chi': 'Χ', r'\\Psi': 'Ψ', r'\\Omega': 'Ω'
        }
        
        # 数学符号映射
        math_symbols = {
            r'\\infty': '∞', r'\\pm': '±', r'\\mp': '∓',
            r'\\times': '×', r'\\div': '÷', r'\\cdot': '·',
            r'\\leq': '≤', r'\\geq': '≥', r'\\neq': '≠',
            r'\\approx': '≈', r'\\equiv': '≡', r'\\sum': '∑',
            r'\\prod': '∏', r'\\int': '∫', r'\\partial': '∂',
            r'\\nabla': '∇', r'\\in': '∈', r'\\notin': '∉',
            r'\\subset': '⊂', r'\\supset': '⊃', r'\\cup': '∪',
            r'\\cap': '∩', r'\\emptyset': '∅'
        }
        
        result = latex_code
        
        # 替换希腊字母
        for latex_symbol, unicode_symbol in greek_letters.items():
            result = re.sub(latex_symbol + r'\b', unicode_symbol, result)
        
        # 替换数学符号
        for latex_symbol, unicode_symbol in math_symbols.items():
            result = re.sub(latex_symbol + r'\b', unicode_symbol, result)
        
        # 处理上标和下标（转换为Unicode上标/下标）
        # 简单的数字上标
        superscript_map = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', 
                          '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹'}
        
        # 简单的数字下标
        subscript_map = {'0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
                        '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉'}
        
        # 处理简单的上标：x^2 → x²
        def replace_superscript(match):
            base = match.group(1)
            sup = match.group(2)
            if sup in superscript_map:
                return base + superscript_map[sup]
            else:
                return f"{base}^{sup}"  # 保持原格式
        
        result = re.sub(r'([a-zA-Z0-9α-ωΑ-Ω])\^{?([0-9])}?', replace_superscript, result)
        
        # 处理简单的下标：x_1 → x₁
        def replace_subscript(match):
            base = match.group(1)
            sub = match.group(2)
            if sub in subscript_map:
                return base + subscript_map[sub]
            else:
                return f"{base}_{sub}"  # 保持原格式
        
        result = re.sub(r'([a-zA-Z0-9α-ωΑ-Ω])_{?([0-9])}?', replace_subscript, result)
        
        # 如果结果与原始输入差别不大，返回None表示无法有效转换
        if len(result) == len(latex_code) and result == latex_code:
            return None
        
        return result
    
    def _insert_math_formula(self, paragraph, latex_code: str, is_display: bool = True) -> bool:
        """
        插入Word原生数学公式
        
        Args:
            paragraph: Word段落对象
            latex_code: LaTeX代码
            is_display: 是否为单行公式
            
        Returns:
            bool: 是否成功插入
        """
        try:
            # 基础LaTeX到OMML的转换
            omml_xml = self._latex_to_omml(latex_code)
            if not omml_xml:
                return False
            
            # 创建数学对象
            math_para = parse_xml(omml_xml)
            
            # 插入到段落中
            paragraph._element.append(math_para)
            
            return True
        except Exception as e:
            print(f"⚠️ 插入数学公式失败: {e}")
            return False
    
    def _latex_to_omml(self, latex_code: str) -> Optional[str]:
        """
        将LaTeX代码转换为Word的OMML格式
        
        Args:
            latex_code: LaTeX代码
            
        Returns:
            str: OMML XML字符串，失败返回None
        """
        try:
            # 基础符号映射
            omml_parts = []
            
            # 分解LaTeX代码
            tokens = self._tokenize_latex(latex_code)
            
            for token in tokens:
                omml_part = self._convert_token_to_omml(token)
                if omml_part:
                    omml_parts.append(omml_part)
            
            if not omml_parts:
                return None
            
            # 组合成完整的OMML
            inner_content = ''.join(omml_parts)
            omml_xml = f'''<m:oMath {nsdecls('m')}>
                <m:r>
                    {inner_content}
                </m:r>
            </m:oMath>'''
            
            return omml_xml
            
        except Exception as e:
            print(f"⚠️ LaTeX到OMML转换失败: {e}")
            return None
    
    def _tokenize_latex(self, latex_code: str) -> List[str]:
        """
        将LaTeX代码分解为token
        
        Args:
            latex_code: LaTeX代码
            
        Returns:
            List[str]: token列表
        """
        tokens = []
        i = 0
        while i < len(latex_code):
            char = latex_code[i]
            
            if char == '\\':
                # LaTeX命令
                if i + 1 < len(latex_code):
                    # 找到完整的命令
                    j = i + 1
                    while j < len(latex_code) and latex_code[j].isalpha():
                        j += 1
                    tokens.append(latex_code[i:j])
                    i = j
                else:
                    i += 1
            elif char in '{}_^':
                # 特殊字符
                tokens.append(char)
                i += 1
            elif char == ' ':
                # 跳过空格
                i += 1
            else:
                # 普通字符，收集连续的字符
                j = i
                while j < len(latex_code) and latex_code[j] not in '\\{}_^ ':
                    j += 1
                if j > i:
                    tokens.append(latex_code[i:j])
                i = j
        
        return tokens
    
    def _convert_token_to_omml(self, token: str) -> Optional[str]:
        """
        将单个token转换为OMML
        
        Args:
            token: LaTeX token
            
        Returns:
            str: OMML片段
        """
        # 希腊字母映射
        greek_map = {
            '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
            '\\epsilon': 'ε', '\\theta': 'θ', '\\lambda': 'λ', '\\mu': 'μ',
            '\\pi': 'π', '\\sigma': 'σ', '\\tau': 'τ', '\\phi': 'φ',
            '\\chi': 'χ', '\\psi': 'ψ', '\\omega': 'ω', '\\nabla': '∇'
        }
        
        # 数学符号映射
        symbol_map = {
            '\\in': '∈', '\\pm': '±', '\\times': '×', '\\cdot': '·',
            '\\leq': '≤', '\\geq': '≥', '\\neq': '≠', '\\approx': '≈',
            '\\sum': '∑', '\\int': '∫', '\\partial': '∂'
        }
        
        if token in greek_map:
            return f'<m:t>{greek_map[token]}</m:t>'
        elif token in symbol_map:
            return f'<m:t>{symbol_map[token]}</m:t>'
        elif token == '\\mathbb':
            # 处理blackboard bold，这里简化处理
            return f'<m:t>𝔹</m:t>'  # 暂时返回一个blackboard字符
        elif token.startswith('\\'):
            # 未知命令，返回原文
            return f'<m:t>{token}</m:t>'
        else:
            # 普通文本
            return f'<m:t>{token}</m:t>'
    
    def _process_markdown_content(self, markdown_content: str, skip_first_title: bool = False):
        """处理Markdown内容"""
        lines = markdown_content.split('\n')
        i = 0
        first_title_skipped = False
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # 处理标题
            if line.startswith('#'):
                # 跳过第一个标题如果需要
                if skip_first_title and not first_title_skipped:
                    first_title_skipped = True
                    i += 1
                    continue
                    
                i = self._process_heading(lines, i)
            
            # 处理代码块
            elif line.startswith('```'):
                i = self._process_code_block(lines, i)
            
            # 处理列表
            elif line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.\s', line):
                i = self._process_list(lines, i)
            
            # 处理图片
            elif line.startswith('!['):
                i = self._process_image(lines, i)
            
            # 处理单行公式
            elif '<<<DISPLAY_FORMULA>>>' in line:
                i = self._process_display_formula_in_doc(lines, i)
            
            # 处理行内公式
            elif '<<<INLINE_FORMULA>>>' in line:
                i = self._process_inline_formula_in_doc(lines, i)
            
            # 处理普通段落
            else:
                i = self._process_paragraph(lines, i)
            
            i += 1
    
    def _process_heading(self, lines: List[str], start_index: int) -> int:
        """处理标题"""
        line = lines[start_index].strip()
        
        # 计算标题级别
        level = 0
        for char in line:
            if char == '#':
                level += 1
            else:
                break
        
        title_text = line[level:].strip()
        if not title_text:
            return start_index
        
        # 根据级别选择样式
        if level == 1:
            # 一级标题处理：创建新节并设置页眉
            if hasattr(self, '_first_h1_added') and self._first_h1_added:
                # 除了第一个一级标题，后续章节创建新节
                section = self.doc.add_section(WD_SECTION.NEW_PAGE)
                # 为新节设置页眉
                self._set_section_header(section, title_text)
                # 验证页眉设置
                self._verify_header_independence(section, title_text)
            else:
                self._first_h1_added = True
                # 为第一章设置页眉（延迟到第一个一级标题出现时）
                # 如果有摘要或目录，创建新节开始第一章
                section = self.doc.add_section(WD_SECTION.NEW_PAGE)
                self._set_section_header(section, title_text)
                # 验证页眉设置
                self._verify_header_independence(section, title_text)
                # 注意：这会让第一章从新页开始，与摘要部分分开
            
            # 一级标题使用自定义样式（居中）
            paragraph = self.doc.add_paragraph(title_text, style='CustomHeading1')
        elif level == 2:
            # 二级标题使用自定义样式（左对齐）
            paragraph = self.doc.add_paragraph(title_text, style='CustomHeading2')
        elif level == 3:
            # 三级标题使用自定义样式（缩进）
            paragraph = self.doc.add_paragraph(title_text, style='CustomHeading3')
        else:
            # 更低级别的标题
            paragraph = self.doc.add_heading(title_text, min(level, 9))
            # 设置字体
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 仅记录当前章节标题；Pandoc 数学转换全局启用
        try:
            self._current_section_title = title_text
            self._use_pandoc_math = True
        except Exception:
            self._use_pandoc_math = True

        # 自然段首行缩进开关逻辑：目录期间关闭；参考文献及以后关闭；其它开启
        try:
            normalized = title_text.strip().lower()
            # 目录：关闭并标记 in_toc
            if any(kw in normalized for kw in ['目录', 'contents', 'table of contents']):
                self._in_toc = True
                self._indent_paragraphs = False
            # 参考文献：关闭缩进（及以后也不缩进）
            elif any(kw in normalized for kw in ['参考文献', 'references', 'bibliography']):
                self._in_toc = False
                self._indent_paragraphs = False
            else:
                # 其他标题：若先前在目录区间，遇到其他标题即退出目录区间并开启缩进
                if self._in_toc:
                    self._in_toc = False
                    self._indent_paragraphs = True
                else:
                    # 非目录、非参考文献：保持开启
                    self._indent_paragraphs = True
        except Exception:
            pass

        return start_index
    
    def _process_code_block(self, lines: List[str], start_index: int) -> int:
        """处理代码块"""
        code_lines = []
        i = start_index + 1
        
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        if code_lines:
            # 添加代码块
            code_paragraph = self.doc.add_paragraph()
            code_run = code_paragraph.add_run('\n'.join(code_lines))
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
            
            # 设置代码块背景和边框
            code_paragraph.paragraph_format.left_indent = Pt(18)
            code_paragraph.paragraph_format.right_indent = Pt(18)
        
        return i
    
    def _process_list(self, lines: List[str], start_index: int) -> int:
        """处理列表（支持多级缩进）"""
        current_line = start_index
        
        while current_line < len(lines):
            line = lines[current_line].strip()
            
            if not line:
                current_line += 1
                continue
                
            # 判断是否还是列表项
            if not (line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.\s', line)):
                break
            
            # 计算缩进级别
            original_line = lines[current_line]
            indent_level = self._calculate_list_indent(original_line)
            
            # 提取列表项文本
            if line.startswith('- ') or line.startswith('* '):
                item_text = line[2:].strip()
            else:
                item_text = re.sub(r'^\d+\.\s*', '', line)
            
            # 根据缩进级别选择样式，并解析行内公式占位符
            if indent_level == 1:
                # 一级列表：自定义样式
                paragraph = self.doc.add_paragraph(style='CustomList1')
                bullet_run = paragraph.add_run('• ')
                bullet_run.font.name = 'Times New Roman'
                bullet_run.font.size = Pt(12)
                bullet_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                self._append_text_with_inline_formulas(paragraph, item_text)
            elif indent_level == 2:
                # 二级列表：缩进18pt
                paragraph = self.doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Pt(36)
                paragraph.paragraph_format.hanging_indent = Pt(18)
                bullet_run = paragraph.add_run('◦ ')
                bullet_run.font.name = 'Times New Roman'
                bullet_run.font.size = Pt(12)
                bullet_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                self._append_text_with_inline_formulas(paragraph, item_text)
            elif indent_level == 3:
                # 三级列表：缩进54pt
                paragraph = self.doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Pt(72)
                paragraph.paragraph_format.hanging_indent = Pt(18)
                bullet_run = paragraph.add_run('▪ ')
                bullet_run.font.name = 'Times New Roman'
                bullet_run.font.size = Pt(12)
                bullet_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                self._append_text_with_inline_formulas(paragraph, item_text)
            else:
                # 更深层级
                paragraph = self.doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Pt(36 * indent_level)
                paragraph.paragraph_format.hanging_indent = Pt(18)
                bullet_run = paragraph.add_run('● ')
                bullet_run.font.name = 'Times New Roman'
                bullet_run.font.size = Pt(12)
                bullet_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                self._append_text_with_inline_formulas(paragraph, item_text)
            
            current_line += 1
        
        return current_line - 1
    
    def _calculate_list_indent(self, line: str) -> int:
        """计算列表项的缩进级别"""
        # 计算前导空格或制表符数量
        indent_chars = 0
        for char in line:
            if char == ' ':
                indent_chars += 1
            elif char == '\t':
                indent_chars += 2  # 制表符算作2个空格
            else:
                break
        
        # 每2个空格算作一级缩进（更符合Markdown约定）
        if indent_chars == 0:
            return 1  # 一级列表
        elif indent_chars <= 2:
            return 2  # 二级列表
        elif indent_chars <= 4:
            return 3  # 三级列表
        else:
            return min(6, (indent_chars // 2) + 1)  # 更深层级，最多6级
    
    def _process_image(self, lines: List[str], start_index: int) -> int:
        """处理图片插入"""
        line = lines[start_index].strip()
        
        # 解析Markdown图片语法: ![alt_text](image_path)
        image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        match = re.match(image_pattern, line)
        
        if match:
            alt_text = match.group(1).strip()  # 图片说明文字
            image_path = match.group(2).strip()  # 图片路径
            
            # 更新统计信息
            self.image_stats['total_found'] += 1
            
            # 在插入图片前，尝试删除紧邻图片上方的“原始图/表标题”段落（章节X-图片Y / 章节X-表格Y），
            # 仅在非“参考图片/参考表格”章节中生效。
            try:
                if not self._is_reference_media_section():
                    self._delete_trailing_media_caption_paragraph()
            except Exception:
                pass

            
            # 尝试插入实际图片
            if self._insert_actual_image(image_path, alt_text):
                self.image_stats['inserted'] += 1
            else:
                self.image_stats['fallback'] += 1
                print(f"⚠️ 图片插入失败，使用降级方案: {alt_text}")
                self._insert_image_fallback(image_path, alt_text)
        else:
            # 如果不是标准图片格式，就作为普通文本处理
            paragraph = self.doc.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置图片引用的格式
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        return start_index
    
    def _validate_image_file(self, image_path: str) -> bool:
        """
        验证图片文件是否存在且有效
        
        Args:
            image_path: 图片文件路径
            
        Returns:
            bool: 文件是否有效
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(image_path):
                print(f"⚠️ 图片文件不存在: {image_path}")
                return False
            
            # 检查文件扩展名
            valid_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
            ext = os.path.splitext(image_path)[1].lower()
            if ext not in valid_extensions:
                print(f"⚠️ 不支持的图片格式: {ext}")
                return False
            
            # 检查文件大小（避免过大文件）
            file_size = os.path.getsize(image_path)
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                print(f"⚠️ 图片文件过大: {file_size / (1024*1024):.1f}MB > 50MB")
                return False
            
            return True
            
        except Exception as e:
            print(f"⚠️ 验证图片文件时出错: {e}")
            return False
    
    def _get_usable_page_width(self) -> float:
        """
        获取页面可用宽度（减去边距）
        
        Returns:
            float: 可用宽度（以inches为单位）
        """
        try:
            # 获取当前节的设置
            if self.doc.sections:
                section = self.doc.sections[-1]  # 使用最后一个节的设置
            else:
                section = self.doc.sections[0]
            
            # 获取页面宽度和边距（转换为inches）
            page_width_inches = section.page_width.inches
            left_margin_inches = section.left_margin.inches
            right_margin_inches = section.right_margin.inches
            
            # 计算可用宽度
            usable_width_inches = page_width_inches - left_margin_inches - right_margin_inches
            
            return usable_width_inches
            
        except Exception as e:
            print(f"⚠️ 获取页面宽度时出错: {e}")
            # 返回默认值（A4纸张大约宽度，减去默认边距）
            return 6.5  # inches
    
    def _calculate_image_size(self, image_path: str, max_width_inches: float) -> tuple:
        """
        计算图片插入尺寸，保持比例不变
        
        Args:
            image_path: 图片路径
            max_width_inches: 最大宽度（inches）
            
        Returns:
            tuple: (width_inches, height_inches) 或 (width_inches, None)
        """
        try:
            if not PIL_AVAILABLE:
                # 没有PIL，使用默认尺寸
                default_width = max_width_inches * 0.8
                return (default_width, None)
            
            # 使用PIL获取图片尺寸
            with Image.open(image_path) as img:
                original_width, original_height = img.size
            
            # 计算比例
            aspect_ratio = original_height / original_width
            
            # 设置目标宽度（页面宽度的90%，留些边距）
            target_width = max_width_inches * 0.9
            target_height = target_width * aspect_ratio
            
            # 限制最大高度（避免图片过高）
            max_height = 8.0  # inches
            if target_height > max_height:
                target_height = max_height
                target_width = target_height / aspect_ratio
            
            return (target_width, target_height)
            
        except Exception as e:
            print(f"⚠️ 计算图片尺寸时出错: {e}")
            # 返回默认尺寸
            default_width = max_width_inches * 0.8
            return (default_width, None)
    
    def _insert_actual_image(self, image_path: str, alt_text: str) -> bool:
        """
        插入实际图片到文档
        
        Args:
            image_path: 图片路径
            alt_text: 图片说明文字
            
        Returns:
            bool: 是否成功
        """
        try:
            # 验证图片文件
            if not self._validate_image_file(image_path):
                return False
            
            # 获取页面可用宽度
            max_width = self._get_usable_page_width()
            
            # 计算合适的尺寸
            width, height = self._calculate_image_size(image_path, max_width)
            
            # 创建段落并插入图片
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run()
            
            # 插入图片
            if height:
                picture = run.add_picture(image_path, width=Inches(width), height=Inches(height))
            else:
                picture = run.add_picture(image_path, width=Inches(width))
            
            # 设置图片居中
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加适当的间距
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            
            # 添加图片说明（如果有）
            if alt_text:
                caption_para = self.doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(alt_text)
                caption_run.font.italic = True
                caption_run.font.size = Pt(10)
                caption_run.font.name = 'Times New Roman'
                caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                # 说明文字的间距
                caption_para.paragraph_format.space_after = Pt(12)
            
            return True
            
        except Exception as e:
            print(f"❌ 插入图片失败: {e}")
            return False
    
    def _insert_image_fallback(self, image_path: str, alt_text: str):
        """
        图片插入失败时的降级处理
        
        Args:
            image_path: 图片路径
            alt_text: 图片说明文字
        """
        try:
            # 创建图片引用段落
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 显示图片信息
            if alt_text:
                info_run = paragraph.add_run(f"[图片: {alt_text}]")
            else:
                info_run = paragraph.add_run(f"[图片]")
            
            info_run.font.italic = True
            info_run.font.size = Pt(11)
            info_run.font.color.rgb = RGBColor(128, 128, 128)
            info_run.font.name = 'Times New Roman'
            info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加路径信息（较小字体）
            path_para = self.doc.add_paragraph()
            path_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            path_run = path_para.add_run(f"路径: {os.path.basename(image_path)}")
            path_run.font.size = Pt(8)
            path_run.font.color.rgb = RGBColor(160, 160, 160)
            path_run.font.name = 'Times New Roman'
            path_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 设置间距
            paragraph.paragraph_format.space_before = Pt(6)
            path_para.paragraph_format.space_after = Pt(12)
            
        except Exception as e:
            print(f"⚠️ 降级处理也失败: {e}")
            # 最简化处理
            simple_para = self.doc.add_paragraph(f"[图片: {alt_text}]" if alt_text else "[图片]")
            simple_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _process_display_formula_in_doc(self, lines: List[str], start_index: int) -> int:
        """
        处理文档中的单行公式标记
        
        Args:
            lines: 文档行列表
            start_index: 当前行索引
            
        Returns:
            int: 处理的行数
        """
        line = lines[start_index].strip()
        
        # 提取公式内容
        formula_match = re.search(r'<<<DISPLAY_FORMULA>>>(.*?)<<<END_DISPLAY_FORMULA>>>', line)
        if formula_match:
            formula_content = formula_match.group(1).strip()
            
            # 创建居中的公式段落
            formula_paragraph = self.doc.add_paragraph()
            formula_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            inserted = False
            # 先做非法/不规范 TeX 预清洗，再做通用归一化
            sanitized, illegal_tex_detected = self._sanitize_illegal_tex(formula_content)
            normalized = self._normalize_tex_in_math(sanitized)

            # 优先使用 Pandoc 数学转换（全局启用）。若 Pandoc 不可用，将在后续文本降级。
            if self._ensure_pandoc():
                result = self._pandoc_formula_to_omml(normalized, is_display=True)
                if result:
                    kind, xml_fragment = result
                    try:
                        if kind == 'oMathPara':
                            # 段落级数学：为 Pandoc 片段补齐命名空间后再解析
                            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            m_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                            wrapper = (
                                f'<w:tmp xmlns:w="{w_ns}" xmlns:m="{m_ns}">' 
                                f'{xml_fragment}'
                                f'</w:tmp>'
                            )
                            tmp_root = parse_xml(wrapper)
                            # 取第一个子节点（应为 m:oMathPara）追加到段落
                            if len(tmp_root):
                                formula_paragraph._element.append(tmp_root[0])
                            else:
                                raise ValueError('oMathPara 解析为空')
                        else:
                            # 行内数学：包入一个 w:r 再追加
                            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            m_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                            wrapper = (
                                f'<w:r xmlns:w="{w_ns}" xmlns:m="{m_ns}">' 
                                f'{xml_fragment}'
                                f'</w:r>'
                            )
                            math_run = parse_xml(wrapper)
                            formula_paragraph._element.append(math_run)
                        inserted = True
                    except Exception as e:
                        print(f"⚠️ 解析/包装 Pandoc OMML 失败，回退: {e}")

            if not inserted:
                # 统一降级：使用 Cambria Math 文本，并追加提示
                formula_run = formula_paragraph.add_run(normalized)
                formula_run.font.name = 'Cambria Math'
                formula_run.font.size = Pt(12)
                formula_run.italic = True
                tip_run = formula_paragraph.add_run(" (非法 Tex，解析公式问题，请直接查看原文)")
                tip_run.font.name = 'Times New Roman'
                tip_run.font.size = Pt(10)
                print(f"⚠️ 发现解析问题，单行公式使用文本显示: {normalized[:10]}...")
            
            # 设置段落格式
            formula_paragraph.paragraph_format.space_before = Pt(12)
            formula_paragraph.paragraph_format.space_after = Pt(12)
            
        
        return start_index  # 返回当前行索引
    
    def _process_inline_formula_in_doc(self, lines: List[str], start_index: int) -> int:
        """
        处理文档中的行内公式标记
        
        Args:
            lines: 文档行列表
            start_index: 当前行索引
            
        Returns:
            int: 处理的行数
        """
        line = lines[start_index].strip()
        
        # 处理包含行内公式的段落
        if '<<<INLINE_FORMULA>>>' in line:
            paragraph = self.doc.add_paragraph()
            paragraph.style = 'Normal'
            # 首行缩进（摘要及以后、参考文献以前生效）
            try:
                if self._indent_paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(self._first_line_indent_pt)
            except Exception:
                pass
            
            # 分割并处理文本和公式
            parts = re.split(r'(<<<INLINE_FORMULA>>>.*?<<<END_INLINE_FORMULA>>>)', line)
            
            for part in parts:
                if part.startswith('<<<INLINE_FORMULA>>>'):
                    # 提取公式内容
                    formula_match = re.search(r'<<<INLINE_FORMULA>>>(.*?)<<<END_INLINE_FORMULA>>>', part)
                    if formula_match:
                        formula_content = formula_match.group(1).strip()
                        sanitized, illegal_tex_detected = self._sanitize_illegal_tex(formula_content)
                        normalized = self._normalize_tex_in_math(sanitized)

                        inserted = False

                        if self._ensure_pandoc():
                            result = self._pandoc_formula_to_omml(normalized, is_display=False)
                            if result:
                                kind, xml_fragment = result
                                try:
                                    if kind == 'oMath':
                                        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                        m_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                                        wrapper = (
                                            f'<w:r xmlns:w="{w_ns}" xmlns:m="{m_ns}">' 
                                            f'{xml_fragment}'
                                            f'</w:r>'
                                        )
                                        math_run = parse_xml(wrapper)
                                        paragraph._element.append(math_run)
                                    else:
                                        # 罕见：行内返回段落级，尝试提取第一个 m:oMath 作为行内插入
                                        inner = self._extract_first_omath(xml_fragment)
                                        if inner:
                                            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                            m_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                                            wrapper = (
                                                f'<w:r xmlns:w="{w_ns}" xmlns:m="{m_ns}">' 
                                                f'{inner}'
                                                f'</w:r>'
                                            )
                                            math_run = parse_xml(wrapper)
                                            paragraph._element.append(math_run)
                                        else:
                                            raise ValueError('无法从 oMathPara 中提取 oMath 用于行内插入')
                                    inserted = True
                                except Exception as e:
                                    print(f"⚠️ 解析/包装 Pandoc OMML 失败，回退: {e}")

                        if not inserted:
                            # 统一降级：使用 Cambria Math 文本，并追加提示
                            formula_run = paragraph.add_run(normalized)
                            formula_run.font.name = 'Cambria Math'
                            formula_run.font.size = Pt(11)
                            formula_run.italic = True
                            tip_run = paragraph.add_run(" (非法 Tex，解析公式问题，请直接查看原文)")
                            tip_run.font.name = 'Times New Roman'
                            tip_run.font.size = Pt(10)
                            print(f"⚠️ 发现解析问题，行内公式使用文本显示: {normalized[:10]}...")
                        
                elif part.strip():
                    # 添加普通文本
                    text_run = paragraph.add_run(part)
                    text_run.font.name = 'Times New Roman'
                    text_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    text_run.font.size = Pt(12)
        
        return start_index  # 返回当前行索引
    
    def _append_text_with_inline_formulas(self, paragraph, text: str) -> None:
        """在给定段落中追加一段可能包含行内公式占位符的文本。
        会将 <<<INLINE_FORMULA>>>...<<<END_INLINE_FORMULA>>> 解析为 OMML 数学，
        其余普通文本通过 _add_formatted_text 添加，从而保留粗体/斜体格式。
        """
        if not text:
            return
        parts = re.split(r'(<<<INLINE_FORMULA>>>.*?<<<END_INLINE_FORMULA>>>)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('<<<INLINE_FORMULA>>>'):
                formula_match = re.search(r'<<<INLINE_FORMULA>>>(.*?)<<<END_INLINE_FORMULA>>>', part)
                if not formula_match:
                    # 不规范片段，按普通文本处理
                    self._add_formatted_text(paragraph, part)
                    continue
                formula_content = formula_match.group(1).strip()
                sanitized, _illegal = self._sanitize_illegal_tex(formula_content)
                normalized = self._normalize_tex_in_math(sanitized)

                inserted = False
                if self._ensure_pandoc():
                    result = self._pandoc_formula_to_omml(normalized, is_display=False)
                    if result:
                        kind, xml_fragment = result
                        try:
                            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            m_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                            if kind == 'oMath':
                                wrapper = (
                                    f'<w:r xmlns:w="{w_ns}" xmlns:m="{m_ns}">' 
                                    f'{xml_fragment}'
                                    f'</w:r>'
                                )
                                math_run = parse_xml(wrapper)
                                paragraph._element.append(math_run)
                            else:
                                # 兼容少见情形：行内返回段落级
                                inner = self._extract_first_omath(xml_fragment)
                                if inner:
                                    wrapper = (
                                        f'<w:r xmlns:w="{w_ns}" xmlns:m="{m_ns}">' 
                                        f'{inner}'
                                        f'</w:r>'
                                    )
                                    math_run = parse_xml(wrapper)
                                    paragraph._element.append(math_run)
                                else:
                                    raise ValueError('无法从 oMathPara 中提取 oMath 用于行内插入')
                            inserted = True
                        except Exception as e:
                            print(f"⚠️ 列表项解析/包装 Pandoc OMML 失败，回退: {e}")

                if not inserted:
                    # 文本降级：Cambria Math 斜体，不追加提示，避免破坏列表排版
                    formula_run = paragraph.add_run(normalized)
                    formula_run.font.name = 'Cambria Math'
                    formula_run.font.size = Pt(11)
                    formula_run.italic = True
            else:
                # 普通文本（含粗体/斜体）
                self._add_formatted_text(paragraph, part)

    def _process_paragraph(self, lines: List[str], start_index: int) -> int:
        """处理普通段落"""
        line = lines[start_index].strip()
        
        # 检查是否包含行内公式
        if '<<<INLINE_FORMULA>>>' in line:
            # 交给专门的行内公式处理函数
            return self._process_inline_formula_in_doc(lines, start_index)
        
        # 创建段落
        paragraph = self.doc.add_paragraph()
        # 若处于“摘要及以后、参考文献以前”，且非列表/代码/图片/公式上下文，则设置首行缩进
        try:
            if self._indent_paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(self._first_line_indent_pt)
        except Exception:
            pass
        
        # 处理文本格式（粗体、斜体等）
        self._add_formatted_text(paragraph, line)
        
        return start_index

    # ---------------- Pandoc 数学集成：辅助方法 ----------------
    def _ensure_pandoc(self) -> bool:
        """确保可用 Pandoc。仅首次检查并记忆结果。"""
        if self._pandoc_ready is not None:
            return self._pandoc_ready
        if not PYPANDOC_AVAILABLE:
            self._pandoc_ready = False
            return False
        try:
            _ = pypandoc.get_pandoc_version()
            self._pandoc_ready = True
        except OSError:
            try:
                pypandoc.download_pandoc()
                self._pandoc_ready = True
            except Exception as e:
                print(f"⚠️ 下载 Pandoc 失败: {e}")
                self._pandoc_ready = False
        return self._pandoc_ready

    def _normalize_tex_in_math(self, math_src: str) -> str:
        """
        数学片段轻量归一化：
        - 规范 _{...} / ^{...} 花括号空格
        - 去除 { ... } 边界空格
        - 合并 \\mathrm/\\operatorname/\\mathbb/\\mathbf/\\boldsymbol/\\mathcal/\\mathsf/\\mathtt 等命令参数内的 PDF 风格空格
        - 温和清理括号/逗号周围空格
        """
        s = math_src
        s = re.sub(r'_\s*\{\s*([^{}]*?)\s*\}', r'_{\1}', s)
        s = re.sub(r'\^\s*\{\s*([^{}]*?)\s*\}', r'^{\1}', s)
        s = re.sub(r'\{\s+', '{', s)
        s = re.sub(r'\s+\}', '}', s)

        cmd_pattern = (
            r'\\('
            r'mathrm|operatorname\*?|mathbf|boldsymbol|mathbb|mathcal|mathsf|mathtt|'
            r'textrm|textbf|textit'
            r')\s*\{\s*([^{}]+?)\s*\}'
        )

        def _join_arg(m: re.Match) -> str:
            cmd = m.group(1)
            arg = m.group(2)
            arg_fixed = re.sub(r'\s+', '', arg)
            return f'\\{cmd}' + '{' + arg_fixed + '}'

        s = re.sub(cmd_pattern, _join_arg, s)
        s = re.sub(r'\(\s+', '(', s)
        s = re.sub(r'\s+\)', ')', s)
        s = re.sub(r'\s+,', ',', s)
        return s

    def _sanitize_illegal_tex(self, math_src: str) -> Tuple[str, bool]:
        """
        针对 Pandoc 易失败的非法/不规范 TeX 进行预清洗：
        - 将 \calX 转为 \mathcal{X}
        - 删除空参数命令（如 \mathrm{} 等）
        - 纠正常见命令的星号位置（如 \operatorname\ast -> \operatorname*）
        返回 (修正后的字符串, 是否检测到不规范)
        """
        text = math_src
        illegal = False

        # \cal X -> \mathcal{X}
        def _cal_to_mathcal(m: re.Match) -> str:
            nonlocal illegal
            illegal = True
            return f"\\mathcal{{{m.group(1)}}}"
        text_new = re.sub(r'\\cal\s*([A-Za-z])', _cal_to_mathcal, text)
        if text_new != text:
            text = text_new

        # { \cal { X } } -> \mathcal{X}
        def _brace_cal_to_mathcal(m: re.Match) -> str:
            nonlocal illegal
            illegal = True
            return f"\\mathcal{{{m.group(1)}}}"
        text = re.sub(r'\{\s*\\cal\s*\{\s*([A-Za-z])\s*\}\s*\}', _brace_cal_to_mathcal, text)

        # 删除空参数命令
        empty_cmd_pattern = r'\\(mathrm|operatorname\*?|mathbf|boldsymbol|mathbb|mathcal|mathsf|mathtt|textrm|textbf|textit)\s*\{\s*\}'
        if re.search(empty_cmd_pattern, text):
            illegal = True
            text = re.sub(empty_cmd_pattern, '', text)

        # \operatorname\ast -> \operatorname*
        if re.search(r'\\operatorname\\ast', text):
            illegal = True
            text = re.sub(r'\\operatorname\\ast', r'\\operatorname*', text)

        # { \bf x } -> \mathbf{x}（Pandoc 不接受 \bf）
        def _bf_to_mathbf(m: re.Match) -> str:
            nonlocal illegal
            illegal = True
            inner = m.group(1)
            inner = re.sub(r'\s+', ' ', inner).strip()
            return f"\\mathbf{{{inner}}}"
        text = re.sub(r'\{\s*\\bf\s*([^{}]*?)\}', _bf_to_mathbf, text)

        return text, illegal

    def _extract_first_omath(self, omml_fragment: str) -> Optional[str]:
        """
        从 oMathPara 片段中提取首个 <m:oMath>...</m:oMath> 子节点。
        """
        m = re.search(r'(\<m:oMath[\s\S]*?\</m:oMath\>)', omml_fragment)
        if m:
            return m.group(1)
        return None

    def _pandoc_formula_to_omml(self, latex_code: str, is_display: bool) -> Optional[Tuple[str, str]]:
        """
        使用 Pandoc 将单个公式转换为 OMML 片段。返回 (kind, xml) 二元组：
        - kind: 'oMathPara' 或 'oMath'
        - xml: 对应元素的原始片段字符串（不附加命名空间声明）
        带缓存（key 为 (latex_code, is_display)）。
        """
        key = (latex_code, is_display)
        if key in self._pandoc_cache:
            return self._pandoc_cache[key]  # type: ignore[return-value]

        if not self._ensure_pandoc():
            self._pandoc_cache[key] = None
            return None

        mini_md = f'$$\n{latex_code}\n$$' if is_display else f'${latex_code}$'

        try:
            with tempfile.TemporaryDirectory() as td:
                out_path = os.path.join(td, 'f.docx')
                pypandoc.convert_text(mini_md, to='docx', format='md', outputfile=out_path)
                with zipfile.ZipFile(out_path, 'r') as zf:
                    xml_bytes = zf.read('word/document.xml')
                xml_text = xml_bytes.decode('utf-8', errors='ignore')

                # 优先整段提取 m:oMathPara（display），否则退回 m:oMath
                para_match = re.search(r'(\<m:oMathPara[\s\S]*?\</m:oMathPara\>)', xml_text)
                if para_match:
                    omml = para_match.group(1)
                    self._pandoc_cache[key] = ('oMathPara', omml)  # type: ignore[assignment]
                    return 'oMathPara', omml

                omath_match = re.search(r'(\<m:oMath[\s\S]*?\</m:oMath\>)', xml_text)
                if omath_match:
                    omml = omath_match.group(1)
                    self._pandoc_cache[key] = ('oMath', omml)  # type: ignore[assignment]
                    return 'oMath', omml
        except Exception as e:
            print(f"⚠️ Pandoc 公式转换失败: {e}")

        self._pandoc_cache[key] = None
        return None
    
    def _add_formatted_text(self, paragraph, text: str):
        """添加格式化文本到段落"""
        # 简化的格式处理，可以后续扩展
        parts = self._split_formatted_text(text)
        
        for part_text, is_bold, is_italic in parts:
            run = paragraph.add_run(part_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = is_bold
            run.font.italic = is_italic
            
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _split_formatted_text(self, text: str) -> List[Tuple[str, bool, bool]]:
        """分割格式化文本（粗体、斜体）"""
        parts = []
        current_pos = 0
        
        # 使用正则表达式匹配粗体和斜体
        # 先处理粗体 **text**
        bold_pattern = r'\*\*([^*]+?)\*\*'
        # 再处理斜体 *text*
        italic_pattern = r'(?<!\*)\*([^*]+?)\*(?!\*)'
        
        # 结合两个模式
        combined_pattern = r'(\*\*[^*]+?\*\*|(?<!\*)\*[^*]+?\*(?!\*))'
        
        matches = list(re.finditer(combined_pattern, text))
        
        for match in matches:
            # 添加匹配前的正常文本
            if match.start() > current_pos:
                normal_text = text[current_pos:match.start()]
                if normal_text:
                    parts.append((normal_text, False, False))
            
            # 处理匹配的格式化文本
            matched_text = match.group(0)
            if matched_text.startswith('**') and matched_text.endswith('**'):
                # 粗体文本
                bold_text = matched_text[2:-2]
                if bold_text:
                    parts.append((bold_text, True, False))
            elif matched_text.startswith('*') and matched_text.endswith('*'):
                # 斜体文本
                italic_text = matched_text[1:-1]
                if italic_text:
                    parts.append((italic_text, False, True))
            
            current_pos = match.end()
        
        # 添加最后的正常文本
        if current_pos < len(text):
            remaining_text = text[current_pos:]
            if remaining_text:
                parts.append((remaining_text, False, False))
        
        # 如果没有找到任何格式化文本，返回原文本
        if not parts:
            parts = [(text, False, False)]
        
        return parts
    
    def _fix_remaining_bold_formatting(self):
        """
        后处理：扫描整个文档，修复所有未处理的**text**格式
        """
        print("🔧 正在修复剩余的加粗格式...")
        
        try:
            fixed_count = 0
            
            # 遍历文档中的所有段落
            for paragraph in self.doc.paragraphs:
                if not paragraph.text:
                    continue
                    
                # 检查段落文本是否包含未处理的**text**格式
                if '**' in paragraph.text:
                    original_text = paragraph.text
                    
                    # 清空段落内容
                    paragraph.clear()
                    
                    # 重新处理并添加格式化文本
                    self._add_formatted_text(paragraph, original_text)
                    fixed_count += 1
            
            # 同样处理表格中的文本（如果有的话）
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if not paragraph.text:
                                continue
                            if '**' in paragraph.text:
                                original_text = paragraph.text
                                paragraph.clear()
                                self._add_formatted_text(paragraph, original_text)
                                fixed_count += 1
            
            if fixed_count > 0:
                print(f"🔧 已修复 {fixed_count} 个段落的加粗格式")
            else:
                print("✅ 没有发现需要修复的加粗格式")
                
        except Exception as e:
            print(f"⚠️ 修复加粗格式时出错: {e}")
    
    def _verify_all_headers(self):
        """
        验证所有章节的页眉设置
        """
        try:
            print(f"🔍 开始验证所有章节的页眉设置...")
            
            sections = self.doc.sections
            
            for i, section in enumerate(sections):
                try:
                    header = section.header
                    is_linked = getattr(header, 'is_linked_to_previous', None)
                    
                    if header.paragraphs and header.paragraphs[0].text:
                        header_text = header.paragraphs[0].text
                    else:
                        print(f"⚠️ 第{i+1}节无页眉内容")
                        
                except Exception as section_error:
                    print(f"⚠️ 检查第{i+1}节时出错: {section_error}")
            
            print(f"✅ 页眉验证完成")
            
        except Exception as e:
            print(f"⚠️ 验证页眉时出错: {e}")

    def _insert_paragraph_after(self, paragraph) -> Paragraph:
        """在指定段落后插入一个新段落，并返回该段落对象。"""
        p = paragraph._p
        new_p = OxmlElement('w:p')
        p.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    def _ensure_section_vertical_center(self, section) -> None:
        """将节的页面垂直对齐设置为居中。"""
        try:
            sectPr = section._sectPr
            vAlign = sectPr.find(qn('w:vAlign'))
            if vAlign is None:
                vAlign = OxmlElement('w:vAlign')
                sectPr.append(vAlign)
            vAlign.set(qn('w:val'), 'center')
        except Exception as e:
            print(f"⚠️ 设置节垂直居中失败: {e}")

    def _format_title_runs(self, paragraph) -> None:
        """统一设置主标题段落的中英文字体、字号和颜色。"""
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            except Exception:
                pass

    def _postprocess_cover_page(self) -> None:
        """封面页后处理：
        - 将首页主标题放置到整页垂直居中且水平居中
        - 统一标题中英文字体为 Times New Roman/宋体，字号 20，加粗，黑色
        - 在标题后插入一行生成时间，右对齐，样式为 Normal
        - 首页不显示页眉（可选：不同首页页眉/页脚）
        """
        if not self.doc.paragraphs:
            return

        # 尝试找到第一个“文档标题”段落（add_heading(..., 0) 生成）
        title_para = None
        for para in self.doc.paragraphs:
            # 以级别 0 的标题、或样式名包含 'Title' 作为主标题判断
            style_name = getattr(para.style, 'name', '') or ''
            if style_name in ('Title',) or (para.text and para.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(para.text) > 0):
                # 进一步通过字号/粗体等特征判断，若需要更严格可再增强
                title_para = para
                break

        if title_para is None:
            # 回退：使用文档首段作为标题
            title_para = self.doc.paragraphs[0]

        # 首页节设置为垂直居中，并不同首页页眉/页脚
        try:
            first_section = self.doc.sections[0]
            self._ensure_section_vertical_center(first_section)
            # 首页页眉与后续不同，尽量减少视觉干扰
            try:
                first_section.different_first_page_header_footer = True
                # 清空首页页眉内容
                if first_section.first_page_header and first_section.first_page_header.paragraphs:
                    for p in first_section.first_page_header.paragraphs:
                        try:
                            p.clear()
                        except Exception:
                            pass
            except Exception:
                pass
        except Exception as e:
            print(f"⚠️ 首页节设置失败: {e}")

        # 标题水平居中 + 统一字体
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._format_title_runs(title_para)

        # 在标题后插入生成时间行，右对齐
        try:
            time_para = self._insert_paragraph_after(title_para)
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
            time_run = time_para.add_run(timestamp)
            # 正文字体与中文设置
            time_run.font.name = 'Times New Roman'
            time_run.font.size = Pt(18)
            try:
                time_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            except Exception:
                pass
            time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception as e:
            print(f"⚠️ 插入生成时间失败: {e}")

    
    def _print_image_statistics(self):
        """
        输出文档处理统计信息（包括HTML清理和图片处理）
        """
        try:

            # 图片处理统计
            img_stats = self.image_stats
            total = img_stats['total_found']
            inserted = img_stats['inserted']
            fallback = img_stats['fallback']
            
            if total == 0:
                print(f"   未发现图片")
            else:
                
                if inserted > 0:
                    success_rate = (inserted / total) * 100
                    print(f"   成功率: {success_rate:.1f}%")
                
                if fallback > 0:
                    print(f"   📝 提示: {fallback} 张图片以文本形式显示")        
            
            print(f"=" * 50)
            print("")  # 空行分隔
            
        except Exception as e:
            print(f"⚠️ 输出统计信息时出错: {e}")

    # ---------------- 图/表原始标题清理：辅助方法 ----------------
    def _looks_like_media_caption(self, text: str) -> bool:
        """判断一段文本是否形如“章节NUM-图片NUM”或“章节NUM-表格NUM”。"""
        s = text.strip()
        if not s:
            return False
        # 统一中文全角/半角连字符
        s_norm = s.replace('—', '-').replace('–', '-')
        # 简单长度门槛，避免误删普通句子
        if len(s_norm) > 40:
            return False
        # 允许末尾存在换行/脚注符等
        s_norm = re.sub(r'[\s\u2028\u2029]+$', '', s_norm)
        # 章节X-图片Y 或 章节X-表格Y
        pattern = r'^章节\s*\d+\s*-\s*(图片|表格)\s*\d+\s*$'
        return re.match(pattern, s_norm) is not None

    def _is_reference_media_section(self) -> bool:
        """判断当前是否处于“参考图片/参考表格”章节下。依据最近的章节标题关键词。"""
        try:
            title = (self._current_section_title or '').strip().lower()
            return any(kw in title for kw in ['参考图片', '参考表格', 'reference figures', 'reference tables'])
        except Exception:
            return False

    def _delete_last_paragraph_matching_text(self, text: str) -> None:
        """删除文档中最后一个与给定文本匹配的非空段落（通常是刚刚写入的上一段原始图/表标题）。"""
        try:
            target = text.strip()
            for para in reversed(self.doc.paragraphs):
                if para.text and para.text.strip() == target:
                    p = para._p
                    parent = p.getparent()
                    if parent is not None:
                        parent.remove(p)
                    break
        except Exception as e:
            print(f"⚠️ 删除原始图/表标题失败: {e}")

    def _delete_trailing_media_caption_paragraph(self) -> None:
        """在当前文档末尾向上回溯，删除紧邻末尾的、看起来像“章节X-图片Y/表格Y”的原始图/表标题段落。
        注意：不依赖 Markdown 的上一行文本，而是直接查看 Word 文档对象模型，
        删除最后一个非空段落中满足 _looks_like_media_caption 的段落（若存在）。
        """
        try:
            for para in reversed(self.doc.paragraphs):
                if not para.text:
                    continue
                text = para.text.strip()
                if not text:
                    continue
                if self._looks_like_media_caption(text):
                    # 为降低误删概率，仅删除“居中对齐”或“长度较短”的段落
                    is_center = (para.alignment == WD_ALIGN_PARAGRAPH.CENTER)
                    if is_center or len(text) <= 40:
                        p = para._p
                        parent = p.getparent()
                        if parent is not None:
                            parent.remove(p)
                        break
                # 一旦遇到非空但不匹配的段落，则停止回溯，避免跨越多个段落误删
                break
        except Exception as e:
            print(f"⚠️ 回溯删除原始图/表标题失败: {e}")


def convert_markdown_to_word(markdown_content: str, word_path: str, title: str = "") -> bool:
    """
    便捷函数：将Markdown内容转换为Word文档
    
    Args:
        markdown_content: Markdown格式内容
        word_path: Word文档保存路径
        title: 文档标题
        
    Returns:
        bool: 转换是否成功
    """
    converter = MarkdownToWordConverter()
    return converter.convert(markdown_content, word_path, title)


if __name__ == "__main__":
    # 直接运行转换指定的Markdown文件
    
    # 指定要转换的Markdown文件路径
    md_file_path = r"D:\Desktop\ZJU\gen_idea\ma_output\_20250812_111713.md"
    
    
    # 检查文件是否存在
    if not os.path.exists(md_file_path):
        print(f"❌ 文件不存在: {md_file_path}")
        print("请检查文件路径是否正确。")
        exit(1)
    
    try:
        # 读取Markdown文件内容
        print(f"📖 正在读取文件: {md_file_path}")
        with open(md_file_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # 生成输出文件路径（与原文件同目录，同名但扩展名为.docx）
        base_name = os.path.splitext(os.path.basename(md_file_path))[0]
        output_dir = os.path.dirname(md_file_path)
        word_file_path = os.path.join(output_dir, f"test{base_name}.docx")
        
        # 从文件名或内容中提取标题
        title = base_name.replace('_', ' ').strip()
        if title.startswith('20'):
            # 如果是时间戳开头，尝试从内容中提取标题
            lines = markdown_content.split('\n')
            for line in lines:
                if line.strip().startswith('# ') and len(line.strip()) > 2:
                    title = line.strip()[2:].strip()
                    break
        
        print(f"📝 文档标题: {title}")
        print(f"📄 输出路径: {word_file_path}")
        
        # 转换为Word文档
        print("🔄 正在转换为Word文档...")
        success = convert_markdown_to_word(markdown_content, word_file_path, title)
        
        if success:
            print("✅ 转换成功！")
            print(f"📁 Word文档已保存到: {word_file_path}")
        else:
            print("❌ 转换失败！")
            
    except Exception as e:
        print(f"❌ 处理文件时出错: {e}")
        print("请检查文件是否正确或文件是否被其他程序占用。")